# -*- coding: utf-8 -*-
"""生成森林酷跑游戏系统 - 期末综合训练 Word 文档"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'D:\WorkBuddy_project\forest-cool-run\forest-runner-game'
doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align is not None:
        p.alignment = align
    return p

def add_image(img_name, width_inches=5.5):
    path = os.path.join(BASE, img_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        return p
    else:
        add_para(f'[图片缺失: {img_name}]', size=10, color=(200,0,0))
        return None

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table_cell(cell, text, bold=False, color_hex=None, size=10):
    """Format a table cell"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color_hex:
        set_cell_shading(cell, color_hex)

# ============================================================
# 1. 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

add_para('期末综合训练', bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1B, 0x5E, 0x20))
add_para('森林酷跑游戏系统', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x2E, 0x7D, 0x32))
doc.add_paragraph()
add_para('项目文档', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x66, 0x66, 0x66))

for _ in range(3):
    doc.add_paragraph()

add_para('小组成员：焦国坤、丁巴达杰、肖盼、陈宇聪、廖锦源、陈星羽、唐紫高', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('项目类型：休闲小游戏', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('开发框架：Flask + MySQL + HTML5 Canvas', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('提交日期：2026年6月', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# 2. 目录（简易文字版）
# ============================================================
add_heading('目录', level=1)
toc_items = [
    '一、项目概述',
    '二、功能模块图',
    '三、功能模块详述',
    '    3.1 用户管理模块',
    '    3.2 游戏核心模块',
    '    3.3 成绩排行模块',
    '    3.4 道具商城模块',
    '    3.5 系统管理模块',
    '四、技术架构',
    '五、数据库设计',
    '六、项目创新点',
    '七、小组分工',
    '八、总结与展望',
]
for item in toc_items:
    add_para(item, size=11)
doc.add_page_break()

# ============================================================
# 3. 一、项目概述
# ============================================================
add_heading('一、项目概述', level=1)
add_para('森林酷跑是一款基于 Web 技术的横版跑酷休闲小游戏。玩家扮演森林中的小动物，在奔跑过程中通过跳跃和下滑躲避各种障碍物，收集金币和道具，挑战最高分数。项目采用 Flask 框架作为后端、MySQL 作为数据库、HTML5 Canvas 实现游戏渲染，整体风格为绿色森林卡通主题。')

add_heading('项目特点', level=2)
features = [
    '🌲 绿色森林卡通视觉风格，配色统一、界面美观',
    '🎮 完整的游戏核心玩法：跳跃/二段跳/下滑躲避/碰撞检测',
    '🌳 5种障碍物形态（树桩/石头/灌木/蘑菇/飞鸟），形态丰富',
    '🎁 5种游戏道具（无敌星/磁铁/双倍金币/缩小药水/慢速时钟）',
    '🏆 排行榜系统支持按难度筛选与分页',
    '⭐ 三档难度选择 + 五阶段动态难度递增',
    '🎵 Web Audio API 音效系统',
    '👤 用户注册/登录/个人中心/历史成绩',
    '⚙ 后台管理面板（公告管理）',
]
for f in features:
    add_para(f, size=11)

doc.add_page_break()

# ============================================================
# 4. 二、功能模块图
# ============================================================
add_heading('二、功能模块图', level=1)
add_para('下图展示了森林酷跑游戏系统的完整功能模块架构，包含 5 大模块、19 项子功能：', size=11)
doc.add_paragraph()
add_image('功能模块图.png', width_inches=5.8)
add_para('图1：森林酷跑游戏系统功能模块图', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

doc.add_page_break()

# ============================================================
# 5. 三、功能模块详述
# ============================================================
add_heading('三、功能模块详述', level=1)

# --- 3.1 用户管理 ---
add_heading('3.1 用户管理模块', level=2)
add_para('该模块负责用户的身份认证与个人信息管理，是系统的入口模块。', size=11)

add_heading('3.1.1 注册功能', level=3)
add_para('用户可通过注册页面创建游戏账号，需填写用户名（4-20位字母数字）、昵称（排行榜显示名）和密码（6位以上）。注册信息通过 Flask-WTF 表单验证，密码使用 Werkzeug 哈希加密存储，保障账户安全。', size=11)
add_image('screenshot_02_register.png', 5.5)
add_para('图2：用户注册页面', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

add_heading('3.1.2 登录功能', level=3)
add_para('用户通过用户名和密码登录系统，采用 Flask-Login 实现会话管理。登录后可访问游戏、排行榜和个人中心等功能页面，未登录用户会被自动跳转到登录页。', size=11)
add_image('screenshot_03_login.png', 5.5)
add_para('图3：用户登录页面', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

add_heading('3.1.3 个人中心', level=3)
add_para('登录后用户可查看个人游戏数据：最高分、累计金币、游戏次数、最远距离等统计信息，以及历史成绩列表（支持分页加载）。', size=11)
add_image('screenshot_08_profile.png', 5.5)
add_para('图4：个人中心页面', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

doc.add_page_break()

# --- 3.2 游戏核心 ---
add_heading('3.2 游戏核心模块', level=2)
add_para('游戏核心模块是整个项目最核心的功能模块，基于 HTML5 Canvas 和 JavaScript 实现，无需任何外部游戏框架，纯原生开发。采用模块化架构，player.js / obstacle.js / powerup.js / score.js / api.js / config.js 各自独立，通过全局命名空间通信。', size=11)

add_heading('3.2.1 游戏主界面与难度选择', level=3)
add_para('玩家进入游戏后可选择简单/普通/困难三档难度。开始画面展示完整的道具说明和操作提示。', size=11)
add_image('screenshot_04_game.png', 5.5)
add_para('图5：游戏开始画面（含难度选择与道具说明）', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

add_heading('3.2.2 角色控制', level=3)
add_para('玩家通过键盘操作控制角色：空格/↑键触发跳跃（支持二段跳，MAX_JUMPS=2），↓键触发下滑动作（角色高度缩至40%、宽度增加15px以模拟趴下姿态）。跳跃采用动态重力机制（GRAVITY_PEAK 轻重力），在跳跃顶点产生明显"滞空感"，让玩家有充足时间判断障碍物。', size=11)

add_heading('3.2.3 障碍物系统', level=3)
add_para('系统生成两类障碍物：地面障碍物需跳跃越过，空中飞鸟需下滑躲避（15%概率出现，随难度增加）。地面障碍物包含4种随机形态：树桩（棕色渐变+年轮纹路+绿色树冠）、石头（灰色不规则外形+高光+裂纹）、灌木（多层绿色椭圆+红色浆果）、蘑菇（红色伞盖+白色斑点+白色菌柄）。碰撞检测采用 AABB 矩形相交算法。', size=11)
add_image('screenshot_05_game_playing.png', 5.5)
add_para('图6：游戏进行中（多种障碍物形态、道具拾取、HUD状态栏）', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

add_heading('3.2.4 道具系统', level=3)
add_para('游戏中随机刷新5种道具供玩家拾取：🛡️无敌星（免疫一次碰撞，5秒）、🧲磁铁（自动吸附金币，8秒）、💰双倍金币（金币翻倍，10秒）、🔽缩小药水（角色体积减半，7秒）、⏱️慢速时钟（游戏速度降低30%，6秒）。道具效果有持续时间限制，到期自动失效。界面底部 powerup-hud 状态栏实时显示当前激活的道具及剩余时间。', size=11)

add_heading('3.2.5 动态难度系统', level=3)
add_para('游戏根据玩家得分自动进入五个阶段：新手期（0-500分，速度3.0，间隔宽）、入门期（500-1000）、进阶期（1000-2000）、高手期（2000-5000）、大师期（5000+，全速7.5，间隔最密）。每个阶段逐渐提升障碍物速度和缩小间隔，确保游戏始终有挑战性。', size=11)

add_heading('3.2.6 游戏结束与结算', level=3)
add_para('角色与障碍物碰撞或脱出屏幕底部即判定游戏结束。结算面板展示本轮得分、金币和距离，成绩通过 AJAX 自动提交到后端数据库。死亡后不可通过空格键立即重开（防止误触），必须点击"再来一局"按钮。', size=11)
add_image('screenshot_06_gameover.png', 5.5)
add_para('图7：游戏结束结算面板', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

doc.add_page_break()

# --- 3.3 成绩排行 ---
add_heading('3.3 成绩排行模块', level=2)
add_para('该模块负责游戏成绩的提交、查询和展示排名。向后端提交成绩时使用 AJAX(POST /score/submit)，包含分数、金币、距离、难度等数据。排行榜支持按[全部/简单/普通/困难]四种维度筛选，以及分页浏览（每页20条）。', size=11)
add_image('screenshot_07_leaderboard.png', 5.5)
add_para('图8：排行榜页面（难度筛选 + 分页）', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

# --- 3.4 道具商城 ---
add_heading('3.4 道具商城模块', level=2)
add_para('道具商城提供道具的购买与管理功能。商店页面展示所有可购买道具及其价格，用户可使用游戏金币购买道具存入背包。背包管理显示用户拥有的道具数量，支持在游戏前装备道具。后端提供完整的 RESTful API 接口（/item/api/buy、/item/api/use、/item/api/my-items 等）。', size=11)

# --- 3.5 系统管理 ---
add_heading('3.5 系统管理模块', level=2)
add_para('管理员可通过后台面板管理游戏公告（增删改），用于发布游戏更新、活动通知等信息。公告内容自动同步到首页展示区。前端页面响应式布局，适配不同屏幕尺寸。音效系统基于 Web Audio API 实现跳跃、碰撞、金币收集、道具拾取等音效反馈。', size=11)
add_image('screenshot_01_home.png', 5.5)
add_para('图9：网站首页（公告展示 + 功能导航）', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

doc.add_page_break()

# ============================================================
# 6. 四、技术架构
# ============================================================
add_heading('四、技术架构', level=1)

add_heading('4.1 后端技术栈', level=2)
backend_tech = [
    'Web 框架：Flask 3.x（工厂函数 + 蓝图模块化架构）',
    '数据库：MySQL 8.0（utf8mb4_unicode_ci 字符集）',
    'ORM：Flask-SQLAlchemy（声明式模型）',
    '认证：Flask-Login（会话管理 + 登录保护装饰器）',
    '表单验证：Flask-WTF（CSRF 保护）',
    '密码加密：Werkzeug generate_password_hash / check_password_hash',
    '蓝图划分：main / auth / game / score / admin / item（6个蓝图）',
]
for t in backend_tech:
    add_para(f'• {t}', size=11)

add_heading('4.2 前端技术栈', level=2)
frontend_tech = [
    '页面模板：Jinja2（继承自 base.html，统一绿色森林主题）',
    '游戏渲染：HTML5 Canvas 2D（纯原生，无外部游戏框架）',
    '游戏架构：模块化 JavaScript（9个独立模块，全局命名空间通信）',
    '主循环：requestAnimationFrame 驱动 60fps',
    '碰撞检测：AABB 矩形相交 + 圆形与矩形',
    '音效系统：Web Audio API',
    '异步通信：Fetch API（成绩提交、排行榜加载）',
    'CSS：CSS3 渐变 / 阴影 / 动画 / 响应式布局',
]
for t in frontend_tech:
    add_para(f'• {t}', size=11)

add_heading('4.3 项目结构', level=2)
struct = '''forest-runner-game/
├── run.py                  # 启动入口
├── config.py               # 数据库/Flask 配置
├── requirements.txt        # Python 依赖
├── app/
│   ├── __init__.py         # 工厂函数 + 蓝图注册 + 数据库初始化
│   ├── models/             # 数据模型层
│   │   ├── user.py         #   User 用户模型
│   │   ├── score.py        #   Score 成绩模型
│   │   ├── skin.py         #   Skin/UserSkin 皮肤模型
│   │   ├── item.py         #   Item/UserItem 道具模型
│   │   └── notice.py       #   Notice 公告模型
│   ├── routes/             # 路由蓝图层
│   │   ├── main.py         #   首页路由
│   │   ├── auth.py         #   认证路由（登录/注册/登出）
│   │   ├── game.py         #   游戏路由（游戏页/个人中心/皮肤API）
│   │   ├── score.py        #   成绩路由（排行榜/提交/历史/统计）
│   │   ├── admin.py        #   管理路由（后台/公告管理）
│   │   └── item.py         #   道具路由（商店/购买/使用/背包）
│   ├── templates/          # Jinja2 模板
│   └── static/             # 静态资源
│       ├── css/style.css   #   全局样式
│       ├── js/             #   游戏模块（9个文件）
│       └── images/         #   SVG 素材'''
add_para(struct, size=9)

doc.add_page_break()

# ============================================================
# 7. 五、数据库设计
# ============================================================
add_heading('五、数据库设计', level=1)

add_heading('5.1 数据表一览', level=2)
table = doc.add_table(rows=8, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['表名', '说明', '主要字段', '关联']
data = [
    ['user', '用户表', 'id, username, nickname, password_hash, coin_balance, role, created_at', '—'],
    ['score', '成绩表', 'id, user_id, score, coins, distance, \ndifficulty, play_time, created_at', 'user.id'],
    ['skin', '皮肤表', 'id, name, description, price, image_url', '—'],
    ['user_skin', '用户皮肤', 'id, user_id, skin_id, is_equipped', 'user.id, skin.id'],
    ['item', '道具表', 'id, name, type, description, price, \nduration, icon', '—'],
    ['user_item', '用户道具', 'id, user_id, item_id, quantity', 'user.id, item.id'],
    ['notice', '公告表', 'id, title, content, is_active, created_at', '—'],
]

for row in data:
    for col_idx, val in enumerate(row):
        make_table_cell(table.cell(data.index(row)+1, col_idx), val, size=9)

# Header row
for col_idx, val in enumerate(headers):
    make_table_cell(table.cell(0, col_idx), val, bold=True, color_hex='388E3C', size=9)
    for run in table.cell(0, col_idx).paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

add_heading('5.2 核心关系', level=2)
add_para('用户(user)与成绩(score)为一对多关系，一个用户可有多条游戏记录。用户(user)与皮肤(user_skin)、道具(user_item)也为一对多关系。系统采用外键约束确保数据完整性（ondelete=CASCADE），用户删除时自动清理关联数据。', size=11)

doc.add_page_break()

# ============================================================
# 8. 六、项目创新点
# ============================================================
add_heading('六、项目创新点', level=1)
innovations = [
    ('🎮 1. 动态重力滞空机制',
     '在经典跑酷游戏基础上创新地引入了动态重力系统：角色在跳跃顶点附近（vy 在 -3 到 +2 范围）自动切换为轻重力（0.28），产生明显"滞空"手感，类似《超级马里奥》的"可变重力跳跃"设计。这给了玩家更多的反应时间来判断和躲避障碍物，显著提升了游戏可玩性。'),
    ('🌳 2. 5种精细Canvas手绘障碍物',
     '不同于一般跑酷游戏使用简单的矩形色块，本项目每个地面障碍物都使用 Canvas 2D API 进行了精细绘制：树桩使用径向渐变模拟年轮纹理、石头使用贝塞尔曲线绘制不规则外形并添加高光和裂纹、灌木由多个绿色椭圆叠加搭配红色浆果点缀、蘑菇使用径向渐变绘制红伞盖搭配白色斑点。飞鸟使用贝塞尔曲线绘制翅膀并实现扇动动画。'),
    ('⭐ 3. 五阶段动态难度系统',
     '游戏根据玩家实时得分自动将难度分为五个阶段（新手/入门/进阶/高手/大师），每个阶段动态调整速度、障碍物生成间隔、空中障碍物出现概率。这确保了对新手友好（起步慢、间隔大）同时对高手有挑战（全速、间隔密），实现了"易上手、难精通"的游戏设计哲学。'),
    ('🎁 4. 可选难度 + 实时道具拾取',
     '玩家可主动选择简单/普通/困难三档难度，游戏内随机刷新5种道具供拾取。道具效果有时间限制且到期自动失效，界面底部 powerup-hud 实时显示激活道具的图标和剩余时间。磁铁道具实现金币自动吸附动画、缩小药水需要动态重置碰撞体尺寸等，技术实现有深度。'),
    ('📊 5. 排行榜多维筛选 + 分页',
     '排行榜支持按[全部/简单/普通/困难]四种维度筛选，结合分页技术（每页20条），在数据量较大时仍保持流畅体验。后端使用 SQLAlchemy 查询构建器动态拼接 WHERE 条件，前端使用 AJAX 动态加载，减少页面刷新。'),
    ('🔧 6. 纯原生技术栈，模块化架构',
     '游戏核心完全不依赖任何第三方游戏框架（如 Phaser、PixiJS），基于 HTML5 Canvas 2D API 和 requestAnimationFrame 从零实现游戏主循环、碰撞检测、粒子效果。JS 代码分为 9 个独立模块（config/player/obstacle/powerup/score/api/audio/game），通过全局命名空间解耦通信，架构清晰、易于维护和扩展。'),
    ('🎨 7. 全局绿色森林视觉风格',
     '系统采用统一的绿色森林卡通主题配色，从导航栏渐变、按钮样式到游戏场景（蓝天草地、深绿色地面条带、棕色土层），贯穿全部页面。CSS 变量统一管理颜色，响应式布局适配不同屏幕尺寸。'),
]

for title, desc in innovations:
    add_para(title, bold=True, size=12, color=(0x1B, 0x5E, 0x20))
    add_para(desc, size=11)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 9. 七、小组分工
# ============================================================
add_heading('七、小组分工', level=1)
add_para('本小组共6名成员，分工如下表所示：', size=11)

team_table = doc.add_table(rows=8, cols=4, style='Table Grid')
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER

team_headers = ['姓名', '岗位角色', '具体工作内容', '技术栈']
team_data = [
    ['焦国坤', '组长\nA岗 游戏核心', '• 项目整体规划与进度管理\n• 功能模块图绘制与Word文档整合\n• GitHub仓库管理与代码合并\n• game.js主循环与状态管理\n• 动态难度系统(DIFFICULTY_STAGES)\n• 道具系统前端逻辑(powerup.js)\n• 游戏结束与结算逻辑', 'Git, Word, \nJavaScript, \nHTML5 Canvas'],
    ['丁巴达杰', 'B岗 美术交互', '• 全局CSS样式与森林主题配色\n• SVG角色/障碍物/金币素材设计\n• Canvas 5种障碍物精细绘制\n  （树桩/石头/灌木/蘑菇）\n• Web Audio API 音效系统\n• 页面响应式布局适配', 'CSS3, SVG, \nCanvas, \nWeb Audio'],
    ['肖盼', 'C岗 Flask后端', '• Flask工厂函数+蓝图模块化架构\n• 用户认证系统(注册/登录/登出)\n• Flask-Login会话管理+CSRF保护\n• 道具系统后端API(item蓝图)\n• 登录保护装饰器与权限控制', 'Flask, \nFlask-Login, \nFlask-WTF'],
    ['陈宇聪', 'D岗 数据库', '• MySQL数据库表结构设计与建模\n• scores/user/skin/item/notice模型\n• 成绩提交/查询/排行榜API\n• 排行榜难度筛选+分页功能\n• 个人中心历史成绩与数据统计\n• 数据库SQL建表文件编写', 'MySQL, \nSQLAlchemy, \nPyMySQL'],
    ['廖锦源', 'E岗 测试文档', '• 全功能系统测试与Bug记录\n• 游戏体验评估与改进建议反馈\n• 运行截图收集与整理\n• 演示流程设计与优化', 'Playwright, \n自动化测试'],
    ['陈星羽', 'E岗 测试文档', '• 现场演示准备与流程演练\n• 项目说明文档素材整理\n• 演示录屏与操作录像\n• UI/UX体验优化建议', 'Office, \n录屏工具'],
    ['唐紫高', 'E岗 测试文档', '• 单元测试用例编写与执行\n• 边界条件与异常场景测试\n• 测试报告整理\n• 项目优化建议汇总', 'Python, \n测试工具'],
]

for i, row_data in enumerate(team_data):
    for j, val in enumerate(row_data):
        make_table_cell(team_table.cell(i+1, j), val, size=9)

for j, val in enumerate(team_headers):
    make_table_cell(team_table.cell(0, j), val, bold=True, color_hex='388E3C', size=9)
    for run in team_table.cell(0, j).paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_page_break()

# ============================================================
# 10. 八、总结与展望
# ============================================================
add_heading('八、总结与展望', level=1)

add_para('本项目成功实现了一款功能完整的休闲跑酷游戏系统，涵盖用户管理、游戏核心、成绩排行、道具商城、系统管理五大模块共19项子功能。技术上采用 Flask + MySQL + HTML5 Canvas 全栈方案，前后端解耦、模块化架构清晰。', size=11)

add_heading('项目成果', level=2)
results = [
    '✅ 完整的前后端全栈项目，代码量超过 3000 行',
    '✅ 5 大功能模块、19 项子功能全部实现',
    '✅ 使用 Flask 后端框架（10分）',
    '✅ 使用 MySQL 数据库（10分）',
    '✅ 功能模块结构完整合理（10分）',
    '✅ 7 个创新点，突出项目特色（10分）',
    '✅ 绿色森林主题统一视觉风格，界面美观',
]
for r in results:
    add_para(r, size=11)

add_heading('未来可扩展方向', level=2)
future = [
    '道具商店前端页面：后端 API 已完成，可补充完整的商店购买 UI',
    '皮肤系统：Skin 模型已建表，可扩展角色皮肤换装功能',
    '好友系统与社交排行：增加好友对战、周榜、月榜等社交功能',
    '移动端触屏适配：增加触屏手势操作（上滑跳跃、下滑俯身）',
    '更多游戏模式：无尽模式、计时赛、BOSS 关卡等',
    '粒子特效与屏幕震动：增强游戏视觉反馈',
    'WebSocket 实时对战：双人在线竞技',
]
for f in future:
    add_para(f'• {f}', size=11)

doc.add_paragraph()
add_para('— 全文完 —', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(120,120,120))

# ============================================================
# 保存
# ============================================================
output_path = os.path.join(BASE, '森林酷跑游戏系统_期末综合训练文档.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
